### Handle file reading and writing for TXT, DOCX, and PDF formats.
### Uses python-docx for DOCX, pdfplumber for reading PDFs, and reportlab for writing PDFs.



from docx import Document
import pdfplumber
from reportlab.pdfgen import canvas
from io import BytesIO

### Txt handling
def read_txt(file):
    return file.read().decode("utf-8")

def write_txt(text):
    return text.encode("utf-8")


### Docx handling
def read_docx(file):
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs])

def write_docx(text):
    doc = Document()
    doc.add_paragraph(text)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


### PDF handling
def read_pdf(file):
    text = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()

            if page_text:
                text.append(page_text)

    return "\n".join(text)

def write_pdf(text):
    buffer = BytesIO()
    c = canvas.Canvas(buffer)

    y = 800
    for line in text.split("\n"):
        c.drawString(50, y, line[:100])
        y -= 15
        if y < 50:
            c.showPage()
            y = 800

    c.save()
    buffer.seek(0)
    return buffer


### Main extraction and export functions
def extract_text(file, file_type):
    if file_type == "txt":
        return read_txt(file)
    elif file_type == "docx":
        return read_docx(file)
    elif file_type == "pdf":
        return read_pdf(file)


def export_file(text, file_type):
    if file_type == "txt":
        return write_txt(text)
    elif file_type == "docx":
        return write_docx(text)
    elif file_type == "pdf":
        return write_pdf(text)